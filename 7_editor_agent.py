# Run it this way: python -m streamlit run 7_editor_agent.py
# Added Editor sub-agent

from anthropic import Anthropic
from dotenv import load_dotenv
from datetime import datetime
import streamlit as st

load_dotenv()

client = Anthropic()

# ============================================================
# AGENT PROMPTS
# ============================================================
research_prompt = """You are a research specialist. Your ONLY job is to research 
a topic thoroughly and produce structured research notes.

You may receive multiple types of source material:
- YouTube video transcripts
- Content from specific web pages
- Uploaded documents (PDFs, text files)
- Preferred sources to prioritize in your web searches

When given a topic:
1. Review ALL provided source material first
2. Use web search to fill gaps and find additional current information
3. When preferred sources are specified, prioritize searching those publications
4. Identify 3-5 key themes and recent developments
5. Note specific facts, statistics, quotes, and sources
6. Explain why this matters to professionals

Output format:
=== RESEARCH NOTES ===
Topic: [topic]
Date: [today's date]

SOURCES REVIEWED:
[List all sources — YouTube videos, web pages, uploaded files, and web searches]

KEY FINDINGS:
[Numbered findings with specific details and sources]

INDUSTRY IMPLICATIONS:
[Why this matters, who it affects, what's changing]

NOTABLE QUOTES/STATS:
[Specific data points worth citing]

Do NOT write articles, posts, or any polished content. 
Your job is ONLY to gather and organize the raw research."""

writer_prompt = """You are a professional content writer. You will receive 
research notes from a research team. Your job is to transform those notes 
into polished content in three formats.

IMPORTANT: Do NOT search the web. Do NOT do your own research. Work ONLY 
from the research notes you are given.

From the research notes, produce:

1. PATREON ARTICLE (400-600 words)
   - Informative but conversational, like explaining to a smart friend
   - Compelling headline
   - Use paragraphs, not bullet points

2. LINKEDIN POST (150-200 words)
   - Professional tone
   - Hook in the first line
   - End with a question or call to action
   - Include 3-5 relevant hashtags

3. INSTAGRAM CAPTION (50-75 words)
   - Casual, punchy, emoji-friendly
   - Include relevant hashtags

Label each section clearly with === PATREON ARTICLE ===, 
=== LINKEDIN POST ===, and === INSTAGRAM CAPTION ===

The industry context may change based on the audience. Default to 
technology and AI trends, but adapt your language and examples if a 
specific industry is mentioned."""

# ============================================================
# AGENT 3: THE EDITOR
# Only job: verify claims against research, tighten writing, catch errors
# ============================================================

editor_prompt = """You are a senior editorial fact-checker and editor. You will receive 
TWO inputs:
1. The original RESEARCH NOTES (the ground truth)
2. The WRITTEN CONTENT produced by a writer

Your job is to edit the written content for quality AND accuracy. Specifically:

FACT-CHECK (MOST IMPORTANT):
- Compare EVERY claim, statistic, quote, and fact in the written content 
  against the research notes
- If a claim appears in the written content but NOT in the research notes, 
  REMOVE IT or flag it clearly with [UNVERIFIED]
- If a statistic or number doesn't match the research notes, correct it
- Do NOT add new information that isn't in the research notes
- Do NOT invent quotes, statistics, or claims to make the content sound better

EDITORIAL QUALITY:
- Tighten language — cut filler words and redundant phrases
- Ensure each format is distinct (Patreon should NOT read like LinkedIn)
- Verify word counts are within range (Patreon: 400-600, LinkedIn: 150-200, 
  Instagram: 50-75)
- Check that the Patreon article uses paragraphs, not bullet points
- Ensure LinkedIn has a strong hook and ends with a CTA
- Ensure Instagram is punchy with relevant hashtags

OUTPUT FORMAT:
First, provide a brief === EDITOR'S NOTES === section listing:
- Any claims you removed or flagged as unverified
- Any factual corrections you made
- Key edits for quality

Then provide the final polished versions:
=== PATREON ARTICLE ===
[edited version]

=== LINKEDIN POST ===
[edited version]

=== INSTAGRAM CAPTION ===
[edited version]

IMPORTANT: Do NOT search the web. Do NOT do your own research. Your ONLY 
source of truth is the research notes provided. If something isn't in the 
research notes, it should NOT be in the final content."""

# ============================================================
# HELPER: Fetch YouTube transcripts
# ============================================================
def get_youtube_transcript(url):
    """Extract transcript text from a YouTube video URL."""
    try:
        from youtube_transcript_api import YouTubeTranscriptApi

        # Extract video ID from various YouTube URL formats
        video_id = None
        if "youtu.be/" in url:
            video_id = url.split("youtu.be/")[1].split("?")[0]
        elif "v=" in url:
            video_id = url.split("v=")[1].split("&")[0]
        elif "shorts/" in url:
            video_id = url.split("shorts/")[1].split("?")[0]

        if not video_id:
            return f"[Could not extract video ID from: {url}]"

        ytt_api = YouTubeTranscriptApi()
        transcript = ytt_api.fetch(video_id)
        text = " ".join([snippet.text for snippet in transcript])
        return text

    except Exception as e:
        return f"[Error getting transcript for {url}: {str(e)}]"


# ============================================================
# HELPER: Fetch webpage content
# ============================================================
def get_webpage_content(url):
    """Extract text content from a webpage URL."""
    try:
        import requests
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.skip_tags = {"script", "style", "nav", "footer", "header"}
                self.current_skip = False

            def handle_starttag(self, tag, attrs):
                if tag in self.skip_tags:
                    self.current_skip = True

            def handle_endtag(self, tag):
                if tag in self.skip_tags:
                    self.current_skip = False

            def handle_data(self, data):
                if not self.current_skip:
                    stripped = data.strip()
                    if stripped:
                        self.text_parts.append(stripped)

        response = requests.get(url, timeout=15, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        response.raise_for_status()

        parser = TextExtractor()
        parser.feed(response.text)
        return " ".join(parser.text_parts)

    except Exception as e:
        return f"[Error reading {url}: {str(e)}]"


# ============================================================
# THE ENGINE: Run any agent with streaming for Streamlit
# ============================================================
def run_agent_stream(system_prompt, user_message, tools=None):
    """Run a single agent. Yields text chunks for streaming."""

    messages = [{"role": "user", "content": user_message}]

    params = {
        "model": "claude-sonnet-4-5-20250929",
        "max_tokens": 4096,
        "system": system_prompt,
        "messages": messages,
    }
    if tools:
        params["tools"] = tools

    with client.messages.stream(**params) as stream:
        for text in stream.text_stream:
            yield text
        response = stream.get_final_message()

    while response.stop_reason != "end_turn":
        messages.append({"role": "assistant", "content": response.content})
        params["messages"] = messages

        with client.messages.stream(**params) as stream:
            for text in stream.text_stream:
                yield text
            response = stream.get_final_message()


# ============================================================
# STREAMLIT INTERFACE
# ============================================================
st.set_page_config(page_title="Research, Content & Editor Agent", layout="wide")
st.title("🔬 Research,  Content & Editor Agent")
st.markdown("Enter a topic and provide sources. The Research Agent will gather information, the Writer Agent will create your content, and the Editor Agent will fact-check and polish the final output.")

# --- INPUT SECTION ---
topic = st.text_input("Enter a topic:", placeholder="e.g., Latest AI agent developments in 2026")

# Two columns for source inputs
col1, col2 = st.columns(2)

with col1:
    youtube_urls = st.text_area(
        "YouTube URLs (one per line)",
        placeholder="https://youtube.com/watch?v=...\nhttps://youtu.be/...",
        height=120
    )

    uploaded_file = st.file_uploader(
        "Upload source material (optional)",
        type=["txt", "md", "csv", "pdf"]
    )

with col2:
    website_urls = st.text_area(
        "Website URLs to read (one per line)",
        placeholder="https://example.com/article\nhttps://news.site.com/story",
        height=120
    )

    preferred_sources = st.text_area(
        "Preferred sources for web search (optional)",
        placeholder="e.g., Skift, MeetingsNet, PCMA, TechCrunch, Ars Technica",
        height=68
    )

# --- RUN BUTTON ---
if st.button("Run Agents", type="primary") and topic:

    # ============================
    # GATHER ALL SOURCE MATERIAL
    # ============================
    source_sections = []

    # 1. YouTube transcripts
    if youtube_urls.strip():
        urls = [u.strip() for u in youtube_urls.strip().split("\n") if u.strip()]
        st.info(f"Fetching {len(urls)} YouTube transcript(s)...")

        for url in urls:
            with st.spinner(f"Reading: {url}"):
                transcript = get_youtube_transcript(url)
                source_sections.append(f"=== YOUTUBE TRANSCRIPT: {url} ===\n{transcript}")

        st.success(f"Loaded {len(urls)} YouTube transcript(s)")

    # 2. Website content
    if website_urls.strip():
        urls = [u.strip() for u in website_urls.strip().split("\n") if u.strip()]
        st.info(f"Reading {len(urls)} webpage(s)...")

        for url in urls:
            with st.spinner(f"Reading: {url}"):
                content = get_webpage_content(url)
                source_sections.append(f"=== WEBPAGE CONTENT: {url} ===\n{content}")

        st.success(f"Loaded {len(urls)} webpage(s)")

    # 3. Uploaded file
    if uploaded_file:
        if uploaded_file.name.endswith(".pdf"):
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            file_content = ""
            for page in pdf_reader.pages:
                file_content += page.extract_text() + "\n"
        else:
            file_content = uploaded_file.read().decode("utf-8")

        source_sections.append(f"=== UPLOADED FILE: {uploaded_file.name} ===\n{file_content}")
        st.success(f"Loaded {len(file_content)} characters from {uploaded_file.name}")

    # 4. Preferred sources guidance
    preferred_guidance = ""
    if preferred_sources.strip():
        preferred_guidance = f"\n\nPREFERRED SOURCES: When searching the web, prioritize these sources: {preferred_sources.strip()}"

    # ============================
    # BUILD RESEARCH INPUT
    # ============================
    research_input = f"Topic: {topic}{preferred_guidance}"

    if source_sections:
        research_input += "\n\nSOURCE MATERIAL PROVIDED:\n\n" + "\n\n".join(source_sections)

    # ============================
    # RUN RESEARCH AGENT
    # ============================
    st.header("📋 Research Agent")
    st.caption("Reviewing sources and searching the web...")

    research_container = st.empty()
    research_text = ""
    for chunk in run_agent_stream(
        system_prompt=research_prompt,
        user_message=research_input,
        tools=[{"type": "web_search_20250305", "name": "web_search", "max_uses": 5}],
    ):
        research_text += chunk
        research_container.markdown(research_text)

    st.success("Research complete!")

    # ============================
    # RUN WRITER AGENT
    # ============================
    st.header("✍️ Writer Agent")
    st.caption("Creating content from research notes...")

    writer_container = st.empty()
    writer_text = ""
    for chunk in run_agent_stream(
        system_prompt=writer_prompt,
        user_message=f"Here are the research notes to work from:\n\n{research_text}",
        tools=None,
    ):
        writer_text += chunk
        writer_container.markdown(writer_text)

    st.success("Content complete!")

    # ============================
    # RUN EDITOR AGENT
    # ============================
    st.header("🔍 Editor Agent")
    st.caption("Fact-checking against research and polishing content...")

    editor_container = st.empty()
    editor_text = ""
    for chunk in run_agent_stream(
        system_prompt=editor_prompt,
        user_message=f"Here are the RESEARCH NOTES (your source of truth):\n\n{research_text}\n\n---\n\nHere is the WRITTEN CONTENT to edit:\n\n{writer_text}",
        tools=None,
    ):
        editor_text += chunk
        editor_container.markdown(editor_text)

    st.success("Editing complete!")

    # ============================
    # SAVE AND DOWNLOAD
    # ============================
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    os.makedirs("Output", exist_ok=True)
    filename = f"Output/output_{timestamp}.docx"

    doc = Document()

    # --- Title ---
    title = doc.add_heading("Research, Content & Editor Agent Output", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Topic: {topic}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph("")  # spacer

    # --- Research Notes ---
    doc.add_heading("Research Notes", level=1)
    doc.add_paragraph(research_text)

    # --- Page break before Writer Draft ---
    doc.add_page_break()
    doc.add_heading("Writer Draft", level=1)
    doc.add_paragraph(writer_text)

    # --- Page break before Editor Final ---
    doc.add_page_break()
    doc.add_heading("Edited Final Version", level=1)
    doc.add_paragraph(editor_text)

    doc.save(filename)

    # Read the file back for download button
    with open(filename, "rb") as f:
        doc_bytes = f.read()

    st.download_button(
        label="📥 Download Output (.docx)",
        data=doc_bytes,
        file_name=f"output_{timestamp}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.info(f"Saved to {filename}")