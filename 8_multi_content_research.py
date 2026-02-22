# Run this way:  python -m streamlit run 8_multi_content_research.py

from anthropic import Anthropic
from dotenv import load_dotenv
from datetime import datetime
import streamlit as st
import yaml
import os

load_dotenv()

# Support both local .env and Streamlit Cloud secrets
if "ANTHROPIC_API_KEY" in st.secrets:
    os.environ["ANTHROPIC_API_KEY"] = st.secrets["ANTHROPIC_API_KEY"]

client = Anthropic()


# ============================================================
# PROFILE LOADER: Read YAML profiles from the profiles/ folder
# ============================================================
def load_profiles(profiles_dir="profiles"):
    """Load all YAML profile files from the profiles directory."""
    profiles = {}
    if not os.path.exists(profiles_dir):
        st.error(f"Profiles directory '{profiles_dir}' not found. Create it and add .yaml files.")
        return profiles

    for filename in sorted(os.listdir(profiles_dir)):
        if filename.endswith((".yaml", ".yml")):
            filepath = os.path.join(profiles_dir, filename)
            with open(filepath, "r", encoding="utf-8") as f:
                profile = yaml.safe_load(f)
                # Use filename (without extension) as the key
                key = os.path.splitext(filename)[0]
                profiles[key] = profile

    return profiles


# ============================================================
# HELPER: Fetch YouTube transcripts
# ============================================================
def get_youtube_transcript(url):
    """Extract transcript text from a YouTube video URL."""
    try:
        from youtube_transcript_api import YouTubeTranscriptApi

        # Extract video ID from various YouTube URL formats
        video_id = None
        if "youtu.be/" in url:
            video_id = url.split("youtu.be/")[1].split("?")[0]
        elif "v=" in url:
            video_id = url.split("v=")[1].split("&")[0]
        elif "shorts/" in url:
            video_id = url.split("shorts/")[1].split("?")[0]

        if not video_id:
            return f"[Could not extract video ID from: {url}]"

        ytt_api = YouTubeTranscriptApi()
        transcript = ytt_api.fetch(video_id)
        text = " ".join([snippet.text for snippet in transcript])
        return text

    except Exception as e:
        return f"[Error getting transcript for {url}: {str(e)}]"


# ============================================================
# HELPER: Fetch webpage content
# ============================================================
def get_webpage_content(url):
    """Extract text content from a webpage URL."""
    try:
        import requests
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.skip_tags = {"script", "style", "nav", "footer", "header"}
                self.current_skip = False

            def handle_starttag(self, tag, attrs):
                if tag in self.skip_tags:
                    self.current_skip = True

            def handle_endtag(self, tag):
                if tag in self.skip_tags:
                    self.current_skip = False

            def handle_data(self, data):
                if not self.current_skip:
                    text = data.strip()
                    if text:
                        self.text_parts.append(text)

        response = requests.get(url, timeout=15, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        response.raise_for_status()

        parser = TextExtractor()
        parser.feed(response.text)
        full_text = " ".join(parser.text_parts)

        # Trim to first ~8000 chars to avoid overwhelming the agent
        if len(full_text) > 8000:
            full_text = full_text[:8000] + "\n\n[Content trimmed — page was very long]"

        return full_text

    except Exception as e:
        return f"[Error reading {url}: {str(e)}]"


# ============================================================
# THE ENGINE: Run any agent with streaming for Streamlit
# ============================================================
def run_agent_stream(system_prompt, user_message, tools=None, usage_tracker=None):
    """Run a single agent. Yields text chunks for streaming.
    If usage_tracker dict is provided, accumulates token counts into it."""

    messages = [{"role": "user", "content": user_message}]

    params = {
        "model": "claude-sonnet-4-5-20250929",
        "max_tokens": 4096,
        "system": system_prompt,
        "messages": messages,
    }
    if tools:
        params["tools"] = tools

    # First streaming call
    with client.messages.stream(**params) as stream:
        for text in stream.text_stream:
            yield text
        response = stream.get_final_message()

    # Track token usage
    if usage_tracker is not None:
        usage_tracker["input_tokens"] += response.usage.input_tokens
        usage_tracker["output_tokens"] += response.usage.output_tokens
        usage_tracker["api_calls"] += 1

    # Agentic loop — keep going if the agent is using tools
    while response.stop_reason != "end_turn":
        messages.append({"role": "assistant", "content": response.content})
        params["messages"] = messages

        with client.messages.stream(**params) as stream:
            for text in stream.text_stream:
                yield text
            response = stream.get_final_message()

        # Track each loop iteration's usage too
        if usage_tracker is not None:
            usage_tracker["input_tokens"] += response.usage.input_tokens
            usage_tracker["output_tokens"] += response.usage.output_tokens
            usage_tracker["api_calls"] += 1


def run_agent_full(system_prompt, user_message, tools=None):
    """Run a single agent. Returns the complete text output (no streaming)."""

    messages = [{"role": "user", "content": user_message}]

    params = {
        "model": "claude-sonnet-4-5-20250929",
        "max_tokens": 4096,
        "system": system_prompt,
        "messages": messages,
    }
    if tools:
        params["tools"] = tools

    with client.messages.stream(**params) as stream:
        response = stream.get_final_message()

    while response.stop_reason != "end_turn":
        messages.append({"role": "assistant", "content": response.content})
        params["messages"] = messages

        with client.messages.stream(**params) as stream:
            response = stream.get_final_message()

    result = ""
    for block in response.content:
        if hasattr(block, "text"):
            result += block.text

    return result


# ============================================================
# STREAMLIT INTERFACE
# ============================================================
st.set_page_config(page_title="Multi-Profile Research Agent", layout="wide")
st.title("🔬 Multi-Profile Research Agent")

# Load all available profiles
profiles = load_profiles()

if not profiles:
    st.error("No profiles found! Add .yaml files to the profiles/ folder.")
    st.stop()

# ============================================================
# SIDEBAR: Profile selection and source inputs
# ============================================================
with st.sidebar:
    st.header("⚙️ Configuration")

    # Profile selector — the key to the whole system
    profile_keys = list(profiles.keys())
    profile_labels = [f"{profiles[k].get('icon', '📄')} {profiles[k]['name']}" for k in profile_keys]

    selected_index = st.selectbox(
        "Select Research Profile:",
        range(len(profile_keys)),
        format_func=lambda i: profile_labels[i],
    )
    selected_key = profile_keys[selected_index]
    profile = profiles[selected_key]

    st.caption(profile.get("description", ""))

    # Show which agents are active for this profile
    st.markdown("---")
    st.markdown("**Active Agents:**")
    st.markdown(f"- 📋 Research Agent (max {profile['research'].get('max_searches', 5)} searches)")
    st.markdown(f"- ✍️ Writer Agent")
    if profile.get("editor", {}).get("enabled", False):
        st.markdown(f"- 🔍 Editor Agent (fact-check enabled)")
    else:
        st.markdown(f"- ~~🔍 Editor Agent~~ (disabled for this profile)")

    st.markdown("---")
    st.markdown("**Display Options:**")
    show_tokens = st.checkbox("Show token usage & cost", value=False)

    st.markdown("---")
    st.markdown("**Source Inputs:**")

    # YouTube URLs
    youtube_urls = st.text_area(
        "YouTube URLs (one per line):",
        height=80,
        placeholder="https://youtube.com/watch?v=...",
    )

    # Website URLs
    website_urls = st.text_area(
        "Website URLs (one per line):",
        height=80,
        placeholder="https://example.com/article",
    )

    # Preferred sources
    preferred_sources = st.text_area(
        "Preferred sources for web search:",
        height=60,
        placeholder="e.g., TechCrunch, Reuters, SEC.gov",
    )

    # File upload
    uploaded_file = st.file_uploader(
        "Upload source material:",
        type=["txt", "md", "csv", "pdf"],
    )


# ============================================================
# MAIN AREA: Topic input and agent execution
# ============================================================

# ============================================================
# HELPER: Build topic string from custom input fields
# ============================================================
def _build_custom_topic(values, fields):
    """Convert structured custom fields into a formatted string for the Research Agent."""
    parts = []
    for field in fields:
        key = field["key"]
        val = values.get(key, "")
        if val and str(val).strip():
            label = field["label"]
            parts.append(f"{label}: {val}")
    return "\n".join(parts)


# ============================================================
# MAIN AREA: Topic input and agent execution
# ============================================================
st.markdown(f"### {profile.get('icon', '📄')} {profile['name']}")
st.markdown(f"*{profile.get('description', '')}*")

# ============================================================
# DYNAMIC INPUT FIELDS: Either custom fields or generic topic box
# ============================================================
custom_inputs = profile.get("custom_inputs", {})
custom_field_values = {}

if custom_inputs.get("enabled", False):
    # Render profile-specific structured input fields
    fields = custom_inputs.get("fields", [])

    # Use two columns for a cleaner layout
    col_left, col_right = st.columns(2)

    for i, field in enumerate(fields):
        # Full width for textareas, alternate columns for everything else
        if field["type"] == "textarea":
            container = st
        else:
            container = col_left if i % 2 == 0 else col_right

        key = field["key"]
        label = field["label"]
        placeholder = field.get("placeholder", "")

        if field["type"] == "text":
            custom_field_values[key] = container.text_input(label, placeholder=placeholder, key=f"custom_{key}")
        elif field["type"] == "number":
            custom_field_values[key] = container.text_input(label, placeholder=placeholder, key=f"custom_{key}")
        elif field["type"] == "textarea":
            custom_field_values[key] = st.text_area(label, placeholder=placeholder, key=f"custom_{key}", height=80)
        elif field["type"] == "select":
            options = field.get("options", [])
            custom_field_values[key] = container.selectbox(label, options, key=f"custom_{key}")

    # Build the topic string from custom fields
    topic = _build_custom_topic(custom_field_values, fields)

    # Check required fields before allowing run
    required_filled = all(
        str(custom_field_values.get(f["key"], "")).strip()
        for f in fields
        if f.get("required", False) and f["type"] != "select"
    )
else:
    # Generic topic input (for Content Creator, Company Research, etc.)
    topic = st.text_area(
        "Enter your topic:",
        placeholder="e.g., Anthropic, AI adoption in hospitality, autonomous vehicle trends...",
    )
    required_filled = bool(topic and topic.strip())

if st.button("🚀 Run Research", type="primary"):
    if not required_filled:
        st.error("Please fill in all required fields before running.")
        st.stop()

    # ============================
    # GATHER ALL SOURCE MATERIAL
    # ============================
    source_sections = []

    # 1. YouTube transcripts
    if youtube_urls.strip():
        urls = [u.strip() for u in youtube_urls.strip().split("\n") if u.strip()]
        st.info(f"Fetching {len(urls)} YouTube transcript(s)...")

        for url in urls:
            with st.spinner(f"Reading: {url}"):
                transcript = get_youtube_transcript(url)
                source_sections.append(f"=== YOUTUBE TRANSCRIPT: {url} ===\n{transcript}")

        st.success(f"Loaded {len(urls)} YouTube transcript(s)")

    # 2. Website content
    if website_urls.strip():
        urls = [u.strip() for u in website_urls.strip().split("\n") if u.strip()]
        st.info(f"Reading {len(urls)} webpage(s)...")

        for url in urls:
            with st.spinner(f"Reading: {url}"):
                content = get_webpage_content(url)
                source_sections.append(f"=== WEBPAGE CONTENT: {url} ===\n{content}")

        st.success(f"Loaded {len(urls)} webpage(s)")

    # 3. Uploaded file
    if uploaded_file:
        if uploaded_file.name.endswith(".pdf"):
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            file_content = ""
            for page in pdf_reader.pages:
                file_content += page.extract_text() + "\n"
        else:
            file_content = uploaded_file.read().decode("utf-8")

        source_sections.append(f"=== UPLOADED FILE: {uploaded_file.name} ===\n{file_content}")
        st.success(f"Loaded {len(file_content)} characters from {uploaded_file.name}")

    # 4. Preferred sources guidance
    preferred_guidance = ""
    if preferred_sources.strip():
        preferred_guidance = f"\n\nPREFERRED SOURCES: When searching the web, prioritize these sources: {preferred_sources.strip()}"

    # ============================
    # INJECT CURRENT DATE INTO PROMPTS
    # ============================
    today = datetime.now().strftime("%B %d, %Y")  # e.g., "February 19, 2026"
    current_year = datetime.now().strftime("%Y")
    date_context = f"\n\nTODAY'S DATE: {today}\n"

    research_system = profile["research"]["system_prompt"] + date_context
    writer_system = profile["writer"]["system_prompt"] + date_context
    editor_system = profile.get("editor", {}).get("system_prompt", "") + date_context

    # ============================
    # BUILD RESEARCH INPUT
    # ============================
    research_input = f"Topic: {topic}{preferred_guidance}"

    if source_sections:
        research_input += "\n\nSOURCE MATERIAL PROVIDED:\n\n" + "\n\n".join(source_sections)

    # ============================
    # INITIALIZE TOKEN TRACKING
    # ============================
    usage = {"input_tokens": 0, "output_tokens": 0, "api_calls": 0}

    # ============================
    # RUN RESEARCH AGENT
    # ============================
    st.header("📋 Research Agent")
    st.caption("Reviewing sources and searching the web...")

    max_searches = profile["research"].get("max_searches", 5)

    research_container = st.empty()
    research_text = ""
    for chunk in run_agent_stream(
        system_prompt=research_system,
        user_message=research_input,
        tools=[{"type": "web_search_20250305", "name": "web_search", "max_uses": max_searches}],
        usage_tracker=usage,
    ):
        research_text += chunk
        research_container.markdown(research_text.replace("$", "\\$"))

    st.success("Research complete!")

    if not research_text.strip():
        research_text = "No research data was returned. Please provide a general overview of the topic."

    # ============================
    # RUN WRITER AGENT
    # ============================
    st.header("✍️ Writer Agent")
    st.caption("Creating content from research notes...")

    writer_container = st.empty()
    writer_text = ""
    for chunk in run_agent_stream(
        system_prompt=writer_system,
        user_message=f"Here are the research notes to work from:\n\n{research_text}",
        tools=None,
        usage_tracker=usage,
    ):
        writer_text += chunk
        writer_container.markdown(writer_text.replace("$", "\\$"))

    st.success("Writing complete!")

    # ============================
    # RUN EDITOR AGENT (if enabled for this profile)
    # ============================
    editor_text = ""
    if profile.get("editor", {}).get("enabled", False):
        st.header("🔍 Editor Agent")
        st.caption("Fact-checking against research notes and polishing...")

        editor_input = (
            f"ORIGINAL RESEARCH NOTES:\n\n{research_text}\n\n"
            f"---\n\n"
            f"WRITTEN CONTENT TO REVIEW:\n\n{writer_text}"
        )

        editor_container = st.empty()
        for chunk in run_agent_stream(
            system_prompt=editor_system,
            user_message=editor_input,
            tools=None,
            usage_tracker=usage,
        ):
            editor_text += chunk
            editor_container.markdown(editor_text.replace("$", "\\$"))

        st.success("Editing complete!")

    # ============================
    # SAVE AND DOWNLOAD
    # ============================
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Create a safe filename from the topic or first custom field
    if custom_inputs.get("enabled", False):
        safe_name = custom_field_values.get("event_name", "") or custom_field_values.get("city_region", "") or "output"
    else:
        safe_name = topic
    safe_name = "".join(c if c.isalnum() or c in " -_" else "" for c in safe_name)[:40].strip()
    base_filename = f"{selected_key}_{safe_name}_{timestamp}"

    # Determine the final content (edited if available, otherwise written)
    final_content = editor_text if editor_text else writer_text

    # Build the full markdown output
    full_output = f"# {profile['name']}: {safe_name}\n"
    full_output += f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n"
    full_output += f"*Profile: {profile['name']}*\n\n"
    full_output += f"---\n\n## RESEARCH NOTES\n\n{research_text}\n\n"
    full_output += f"---\n\n## WRITTEN CONTENT\n\n{writer_text}\n\n"

    if editor_text:
        full_output += f"---\n\n## EDITED CONTENT\n\n{editor_text}\n\n"
        
    os.makedirs("Output", exist_ok=True)
    md_filename = f"Output/{base_filename}.md"
    with open(md_filename, "w", encoding="utf-8") as f:
        f.write(full_output)

    # --- Generate Word Document ---
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()

        # Title
        title = doc.add_heading(f"{profile['name']}: {safe_name}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Profile: {profile['name']}")
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True

        doc.add_paragraph("")  # spacer

        # Helper: parse markdown-ish text into Word paragraphs
        def add_content_to_doc(doc, text, section_title=None):
            if section_title:
                doc.add_heading(section_title, level=1)

            for line in text.split("\n"):
                line = line.rstrip()

                # Skip empty lines
                if not line:
                    doc.add_paragraph("")
                    continue

                # Detect heading patterns: === HEADING === or ## Heading
                if line.startswith("===") and line.endswith("==="):
                    heading_text = line.strip("= ").strip()
                    doc.add_heading(heading_text, level=2)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith("--- ") or line == "---":
                    # horizontal rule — just add a spacer
                    doc.add_paragraph("")
                elif re.match(r"^\d+\.\s", line):
                    # Numbered item
                    doc.add_paragraph(line, style="List Number")
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    # Regular paragraph — handle bold markers
                    p = doc.add_paragraph()
                    # Simple bold handling: split on ** markers
                    parts = line.split("**")
                    for idx, part in enumerate(parts):
                        if part:
                            run = p.add_run(part)
                            if idx % 2 == 1:  # odd parts are between ** markers
                                run.bold = True

        # Add the final report content (edited version if available)
        add_content_to_doc(doc, final_content, section_title=None)

        # Save Word doc
        docx_filename = f"Output/{base_filename}.docx"
        doc.save(docx_filename)

        # Read the docx file for download
        with open(docx_filename, "rb") as f:
            docx_bytes = f.read()

        word_available = True

    except ImportError:
        word_available = False
        st.warning("Install python-docx for Word document output: `pip install python-docx`")

    # Download buttons
    st.markdown("---")

    dl_col1, dl_col2 = st.columns(2)

    if word_available:
        dl_col1.download_button(
            label="📥 Download Word Document",
            data=docx_bytes,
            file_name=f"{base_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    dl_col2.download_button(
        label="📥 Download Markdown",
        data=full_output,
        file_name=md_filename,
        mime="text/markdown",
    )

    st.info(f"Saved locally to {md_filename}" + (f" and {docx_filename}" if word_available else ""))

    # ============================
    # TOKEN USAGE SUMMARY (if enabled)
    # ============================
    if show_tokens:
        st.markdown("---")
        st.subheader("📊 Token Usage Summary")

        total_tokens = usage["input_tokens"] + usage["output_tokens"]

        # Sonnet pricing: $3 per million input, $15 per million output
        input_cost = (usage["input_tokens"] / 1_000_000) * 3.00
        output_cost = (usage["output_tokens"] / 1_000_000) * 15.00
        total_cost = input_cost + output_cost

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Input Tokens", f"{usage['input_tokens']:,}")
        col2.metric("Output Tokens", f"{usage['output_tokens']:,}")
        col3.metric("Total Tokens", f"{total_tokens:,}")
        col4.metric("Estimated Cost", f"${total_cost:.4f}")

        st.caption(
            f"API calls: {usage['api_calls']}  |  "
            f"Model: claude-sonnet-4-5  |  "
            f"Pricing: USD 3.00/M input, USD 15.00/M output"
        )